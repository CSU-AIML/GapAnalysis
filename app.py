import streamlit as st
import pandas as pd
import numpy as np
import torch
import torch.nn as nn
from transformers import AutoModelForSequenceClassification, AutoTokenizer, AutoModelForSeq2SeqLM, RobertaModel, RobertaTokenizer, RobertaConfig
import io
import os
from huggingface_hub import login, hf_hub_download
from safetensors.torch import load_file
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots

# Additional imports for PDF and DOCX support
import PyPDF2
import pdfplumber
from docx import Document
import zipfile

# Configure Streamlit page
st.set_page_config(
    page_title="Compliance Maturity & Gap Analysis Platform",
    page_icon="🎯",
    layout="wide"
)

# Custom CSS for better styling
st.markdown("""
<style>
.metric-card {
    background-color: #f0f2f6;
    padding: 1rem;
    border-radius: 0.5rem;
    border-left: 4px solid #ff6b6b;
}
.gap-analysis-section {
    background-color: #e8f4fd;
    padding: 1.5rem;
    border-radius: 0.5rem;
    margin: 1rem 0;
}
.recommendation-box {
    background-color: #d4edda;
    padding: 1rem;
    border-radius: 0.5rem;
    border-left: 4px solid #28a745;
}
.gap-form {
    background-color: #f8f9fa;
    padding: 1.5rem;
    border-radius: 0.5rem;
    border: 1px solid #dee2e6;
    margin: 1rem 0;
}
.file-preview {
    background-color: #f8f9fa;
    padding: 1rem;
    border-radius: 0.5rem;
    border: 1px solid #e9ecef;
    max-height: 400px;
    overflow-y: auto;
    font-family: monospace;
    font-size: 0.85em;
}
.evidence-recommendation {
    background-color: #fff3cd;
    padding: 1rem;
    border-radius: 0.5rem;
    border-left: 4px solid #ffc107;
    margin: 1rem 0;
}
.evidence-item {
    background-color: #000000;
    color: #ffffff;
    padding: 0.5rem;
    margin: 0.3rem 0;
    border-radius: 0.3rem;
    border-left: 3px solid #007bff;
}
</style>
""", unsafe_allow_html=True)

# Configuration - Use environment variables for Render deployment
YOUR_HF_ORGANIZATION = os.getenv("HF_ORGANIZATION")
HF_TOKEN = os.getenv("HF_TOKEN")

# Validate configuration
if not YOUR_HF_ORGANIZATION or not HF_TOKEN:
    st.error("❌ Missing required environment variables.")
    st.info("Please ensure HF_ORGANIZATION and HF_TOKEN are set in your Render environment variables.")
    st.info("Go to your Render dashboard > Service Settings > Environment Variables to add them.")
    st.stop()

# Model paths from your Hugging Face organization
MATURITY_MODEL_PATH = f"{YOUR_HF_ORGANIZATION}/compliance-maturity-classifier"
GAP_ANALYSIS_MODEL_PATH = f"{YOUR_HF_ORGANIZATION}/flan-t5-large-gap-analysis"
EVIDENCE_RECOMMENDATION_MODEL_PATH = f"{YOUR_HF_ORGANIZATION}/evidence-recommendation-model"

# Evidence Recommendation Model Classes
class EvidenceRecommendationModel(nn.Module):
    """Custom Evidence Recommendation Model that works with the deployed weights."""
    
    def __init__(self, num_labels=10, hidden_dim=256, dropout_rate=0.3):
        super().__init__()
        self.num_labels = num_labels
        
        # Load RoBERTa config and model
        roberta_config = RobertaConfig.from_pretrained('roberta-base')
        self.roberta = RobertaModel(roberta_config)
        
        # Classification head (same as training)
        self.dropout = nn.Dropout(dropout_rate)
        self.classifier = nn.Sequential(
            nn.Linear(roberta_config.hidden_size, hidden_dim),
            nn.ReLU(),
            nn.Dropout(dropout_rate),
            nn.Linear(hidden_dim, hidden_dim // 2),
            nn.ReLU(),
            nn.Dropout(dropout_rate),
            nn.Linear(hidden_dim // 2, num_labels)
        )
    
    def forward(self, input_ids, attention_mask):
        outputs = self.roberta(input_ids=input_ids, attention_mask=attention_mask)
        pooled_output = outputs.last_hidden_state[:, 0, :]
        pooled_output = self.dropout(pooled_output)
        logits = self.classifier(pooled_output)
        return logits

# Load evidence recommendation model from Hugging Face
@st.cache_resource
def load_evidence_recommendation_model():
    try:
        # Login to Hugging Face (for private models)
        login(token=HF_TOKEN)

        print(f"📥 Loading evidence recommendation model from {EVIDENCE_RECOMMENDATION_MODEL_PATH}...")
        
        # Evidence types (same as training)
        evidence_types = [
            "Implementation Evidence",
            "Governance Documentation", 
            "Procedural Documents",
            "Operational Records",
            "Audit and Assessment Reports",
            "Training and Awareness",
            "Agreements and Legal Documents",
            "Communications and Meeting Records",
            "Planning Documents",
            "Evidence of Continuous Improvement"
        ]
        
        # Load tokenizer
        tokenizer = RobertaTokenizer.from_pretrained(EVIDENCE_RECOMMENDATION_MODEL_PATH)
        
        # Initialize model
        model = EvidenceRecommendationModel(
            num_labels=len(evidence_types),
            hidden_dim=256,
            dropout_rate=0.3
        )
        
        # Load model weights
        model_file = hf_hub_download(
            repo_id=EVIDENCE_RECOMMENDATION_MODEL_PATH, 
            filename="model.safetensors"
        )
        
        # Load state dict
        state_dict = load_file(model_file)
        
        # Load with strict=False to handle missing dropout weights
        missing_keys, unexpected_keys = model.load_state_dict(state_dict, strict=False)
        
        model.eval()
        
        print(f"✅ Evidence recommendation model loaded successfully!")
        return model, tokenizer, evidence_types
        
    except Exception as e:
        st.error(f"Error loading evidence recommendation model: {str(e)}")
        return None, None, None

# Load maturity classification model from Hugging Face
@st.cache_resource
def load_maturity_model():
    try:
        # Login to Hugging Face (for private models)
        login(token=HF_TOKEN)

        print(f"📥 Loading maturity model from {MATURITY_MODEL_PATH}...")
        tokenizer = AutoTokenizer.from_pretrained(MATURITY_MODEL_PATH)
        model = AutoModelForSequenceClassification.from_pretrained(MATURITY_MODEL_PATH)
        model.eval()
        return model, tokenizer
    except Exception as e:
        st.error(f"Error loading maturity model: {str(e)}")
        return None, None

# Load gap analysis model from Hugging Face
@st.cache_resource
def load_gap_analysis_model():
    try:
        # Login to Hugging Face (for private models)
        login(token=HF_TOKEN)

        print(f"📥 Loading gap analysis model from {GAP_ANALYSIS_MODEL_PATH}...")
        tokenizer = AutoTokenizer.from_pretrained(GAP_ANALYSIS_MODEL_PATH)
        model = AutoModelForSeq2SeqLM.from_pretrained(
            GAP_ANALYSIS_MODEL_PATH,
            torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
            device_map="auto" if torch.cuda.is_available() else None
        )
        return model, tokenizer, GAP_ANALYSIS_MODEL_PATH
    except Exception as e:
        print(f"❌ Failed to load gap analysis model: {e}")
        st.error(f"Error loading gap analysis model: {str(e)}")
        return None, None, None

# Define maturity levels
MATURITY_LEVELS = {
    0: "Level 0 - No Evidence",
    1: "Level 1 - Initial",
    2: "Level 2 - Developing",
    3: "Level 3 - Defined",
    4: "Level 4 - Managed and Measurable",
    5: "Level 5 - Optimized"
}

def predict_evidence_types(model, tokenizer, evidence_types, question, domain, framework, threshold=0.5, top_k=5):
    """Predict evidence types for a given question."""
    
    # Format input text (same as training)
    text = f"{question} [SEP] {domain}: {framework}"
    
    # Tokenize
    inputs = tokenizer(
        text, 
        return_tensors="pt", 
        truncation=True, 
        padding=True, 
        max_length=512
    )
    
    # Predict
    with torch.no_grad():
        logits = model(inputs['input_ids'], inputs['attention_mask'])
        probabilities = torch.sigmoid(logits).cpu().numpy()[0]
    
    # Create results
    results = []
    for i, evidence_type in enumerate(evidence_types):
        results.append({
            'evidence_type': evidence_type,
            'probability': float(probabilities[i]),
            'recommended': probabilities[i] > threshold
        })
    
    # Sort by probability (highest first)
    results.sort(key=lambda x: x['probability'], reverse=True)
    
    # Apply top-k
    recommended_evidence = results[:top_k]
    
    return {
        'all_predictions': results,
        'recommended_evidence': recommended_evidence,
        'summary': {
            'total_recommended': len([r for r in results if r['recommended']]),
            'max_probability': float(np.max(probabilities)),
            'avg_probability': float(np.mean(probabilities)),
            'threshold_used': threshold
        }
    }

def display_evidence_recommendations(recommendations, question):
    """Display evidence recommendations in a nice format."""
    
    st.markdown('<div class="evidence-recommendation">', unsafe_allow_html=True)
    st.markdown("### 💡 Recommended Evidence Types")
    st.markdown(f"**For question:** {question}")
    
    # Summary metrics
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Top Recommendation", 
                 f"{recommendations['recommended_evidence'][0]['probability']:.1%}",
                 help="Highest confidence evidence type")
    with col2:
        st.metric("Threshold Met", 
                 recommendations['summary']['total_recommended'],
                 help="Evidence types above 50% confidence")
    with col3:
        st.metric("Average Confidence", 
                 f"{recommendations['summary']['avg_probability']:.1%}",
                 help="Average across all evidence types")
    
    st.markdown("---")
    
    # Top recommendations
    st.markdown("#### 🎯 Top 5 Evidence Types:")
    
    for i, rec in enumerate(recommendations['recommended_evidence'], 1):
        confidence = rec['probability']
        evidence_type = rec['evidence_type']
        
        # Color coding based on confidence
        if confidence >= 0.7:
            icon = "🟢"
            confidence_level = "High"
        elif confidence >= 0.5:
            icon = "🟡"
            confidence_level = "Medium"
        else:
            icon = "🔵"
            confidence_level = "Low"
        
        st.markdown(f"""
        <div class="evidence-item">
            <strong>{icon} {i}. {evidence_type}</strong><br>
            <small>Confidence: {confidence:.1%} ({confidence_level})</small>
        </div>
        """, unsafe_allow_html=True)
    
    # Show all predictions in expandable section
    with st.expander("📊 View All Evidence Types (with probabilities)"):
        df_results = pd.DataFrame(recommendations['all_predictions'])
        df_results['probability'] = df_results['probability'].apply(lambda x: f"{x:.1%}")
        df_results['recommended'] = df_results['recommended'].apply(lambda x: "✅" if x else "❌")
        
        st.dataframe(
            df_results[['evidence_type', 'probability', 'recommended']],
            column_config={
                'evidence_type': 'Evidence Type',
                'probability': 'Confidence',
                'recommended': 'Recommended'
            },
            hide_index=True,
            use_container_width=True
        )
    
    st.markdown('</div>', unsafe_allow_html=True)

def extract_text_from_pdf(uploaded_file):
    """Extract text from PDF file using multiple methods"""
    try:
        # Method 1: Try pdfplumber first (better for complex PDFs)
        try:
            pdf_bytes = uploaded_file.read()
            text = ""

            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            if text.strip():
                return text
        except Exception as e:
            print(f"pdfplumber failed: {e}")

        # Method 2: Fallback to PyPDF2
        try:
            uploaded_file.seek(0)  # Reset file pointer
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""

            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            if text.strip():
                return text
        except Exception as e:
            print(f"PyPDF2 failed: {e}")

        return "Error: Could not extract text from PDF. The PDF might be image-based or corrupted."

    except Exception as e:
        return f"Error processing PDF: {str(e)}"

def extract_text_from_docx(uploaded_file):
    """Extract text from DOCX file"""
    try:
        # Read the uploaded file
        docx_bytes = uploaded_file.read()

        # Create a Document object from bytes
        doc = Document(io.BytesIO(docx_bytes))

        # Extract text from paragraphs
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        if not text.strip():
            return "Warning: No text found in the DOCX file."

        return text

    except Exception as e:
        return f"Error processing DOCX file: {str(e)}"

def process_uploaded_file(uploaded_file):
    """Process uploaded file and extract text content with support for multiple formats"""
    try:
        file_type = uploaded_file.type
        file_name = uploaded_file.name.lower()

        # Text files
        if file_type == "text/plain" or file_name.endswith('.txt'):
            stringio = io.StringIO(uploaded_file.getvalue().decode("utf-8"))
            content = stringio.read()

        # Excel files
        elif file_type in ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"] or file_name.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(uploaded_file)
            content = df.to_string()

        # CSV files
        elif file_type == "text/csv" or file_name.endswith('.csv'):
            df = pd.read_csv(uploaded_file)
            content = df.to_string()

        # PDF files
        elif file_type == "application/pdf" or file_name.endswith('.pdf'):
            content = extract_text_from_pdf(uploaded_file)

        # DOCX files
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_name.endswith('.docx'):
            content = extract_text_from_docx(uploaded_file)

        else:
            st.error(f"Unsupported file type: {file_type}. Please upload TXT, CSV, XLSX, PDF, or DOCX files.")
            return None

        return content

    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return None

def prepare_maturity_input_text(question, evidence_text):
    """Prepare input text for maturity model"""
    formatted_text = f"Question: {question.strip()} Evidence: {evidence_text.strip()}"
    return formatted_text

def predict_maturity(model, tokenizer, text):
    """Predict the current maturity level"""
    inputs = tokenizer(
        text,
        return_tensors="pt",
        padding="max_length",
        truncation=True,
        max_length=512
    )

    with torch.no_grad():
        outputs = model(**inputs)
        predictions = torch.nn.functional.softmax(outputs.logits, dim=-1)
        predicted_class = torch.argmax(predictions, dim=-1).item()
        confidence = predictions[0][predicted_class].item()

    return predicted_class, confidence, predictions[0].tolist()

def create_gap_analysis_prompt(question, evidence, current_level, target_level, domain, framework):
    """Create focused prompt for gap analysis"""

    domain_context = {
        "Banking": "regulatory compliance and financial data protection",
        "Healthcare": "patient data privacy and HIPAA compliance",
        "Technology": "cybersecurity and data protection",
        "Manufacturing": "operational security and safety compliance"
    }.get(domain, "regulatory compliance")

    framework_context = {
        "RBI IT Framework": "Reserve Bank of India IT governance requirements",
        "HIPAA": "Health Insurance Portability and Accountability Act standards",
        "ISO 27001": "Information Security Management System standards",
        "SOX": "Sarbanes-Oxley financial reporting controls"
    }.get(framework, "compliance standards")

    prompt = f"""As a compliance expert specializing in {domain_context}, analyze this gap for {framework_context}:

SITUATION:
Question: {question}
Current State: {evidence}
Maturity: Level {current_level} → Level {target_level}
Context: {domain} | {framework}

Provide a complete gap analysis with these components:

1. Gap Description: What specific controls, processes, or documentation are missing or inadequate?

2. Gap Initiative: What are the top 3 specific actions needed to close this gap (use secure words)?

3. Common Gap Description: What is the general category of this compliance gap?

4. Severity Assessment: What is the risk level (High/Medium/Low) and why?

Complete Gap Analysis:"""

    return prompt

def generate_gap_analysis(gap_model, gap_tokenizer, prompt, model_name, question, evidence, current_level, target_level, domain, framework):
    """Generate gap analysis using the FLAN-T5 model"""

    inputs = gap_tokenizer(prompt, max_length=1024, truncation=True, return_tensors="pt")

    if torch.cuda.is_available() and hasattr(gap_model, 'device'):
        inputs = {k: v.to(gap_model.device) for k, v in inputs.items()}

    generation_config = {
        "max_length": 600,
        "min_length": 100,
        "temperature": 0.8,
        "do_sample": True,
        "top_p": 0.9,
        "repetition_penalty": 1.4,
        "length_penalty": 1.2,
        "num_beams": 3,
        "early_stopping": True,
        "pad_token_id": gap_tokenizer.pad_token_id,
        "eos_token_id": gap_tokenizer.eos_token_id
    }

    try:
        with torch.no_grad():
            outputs = gap_model.generate(**inputs, **generation_config)

        result = gap_tokenizer.decode(outputs[0], skip_special_tokens=True)

        if prompt in result:
            result = result.replace(prompt, "").strip()

        # Clean up incomplete sentences
        if not result.endswith('.'):
            sentences = result.split('.')
            if len(sentences) > 1:
                result = '.'.join(sentences[:-1]) + '.'

        return result

    except Exception as e:
        # Use the passed parameters for fallback
        return create_rule_based_gap_analysis(question, evidence, current_level, target_level, domain, framework)

def create_rule_based_gap_analysis(question, evidence, current_level, target_level, domain, framework):
    """Fallback rule-based gap analysis"""

    evidence_lower = evidence.lower()
    maturity_gap = target_level - current_level

    # Identify gap type
    if any(keyword in evidence_lower for keyword in ['access', 'provisioning', 'user']):
        gap_desc = f"Current access management framework lacks comprehensive user lifecycle procedures and role-based access controls required for Level {target_level} maturity."
        gap_initiative = "Implement automated IAM system with RBAC, develop user lifecycle procedures, establish access review processes"
        common_gap = "Access Management Control Gap"
    elif any(keyword in evidence_lower for keyword in ['incident', 'security', 'tracking']):
        gap_desc = f"Current incident management relies on informal processes without centralized tracking, escalation, or compliance reporting capabilities."
        gap_initiative = "Deploy incident management system, establish formal procedures, implement compliance reporting"
        common_gap = "Incident Management Process Gap"
    elif any(keyword in evidence_lower for keyword in ['data', 'backup', 'recovery']):
        gap_desc = f"Current data management practices lack systematic backup, recovery, and data governance controls required for Level {target_level} compliance."
        gap_initiative = "Implement data governance framework, establish backup and recovery procedures, deploy data classification controls"
        common_gap = "Data Management Control Gap"
    elif any(keyword in evidence_lower for keyword in ['monitoring', 'logging', 'audit']):
        gap_desc = f"Current monitoring and logging capabilities are insufficient for comprehensive audit trails and real-time threat detection required for Level {target_level}."
        gap_initiative = "Deploy centralized logging system, implement continuous monitoring, establish audit trail procedures"
        common_gap = "Monitoring and Logging Gap"
    elif any(keyword in evidence_lower for keyword in ['training', 'awareness', 'education']):
        gap_desc = f"Current security awareness program lacks structured training, regular assessments, and role-based education required for Level {target_level} maturity."
        gap_initiative = "Develop comprehensive training program, implement regular awareness campaigns, establish role-based security education"
        common_gap = "Security Awareness Training Gap"
    elif any(keyword in evidence_lower for keyword in ['risk', 'assessment', 'management']):
        gap_desc = f"Current risk management framework lacks formal assessment procedures, quantitative analysis, and continuous monitoring required for Level {target_level}."
        gap_initiative = "Implement risk assessment framework, establish quantitative risk analysis, deploy continuous risk monitoring"
        common_gap = "Risk Management Framework Gap"
    elif any(keyword in evidence_lower for keyword in ['policy', 'procedure', 'documentation']):
        gap_desc = f"Current policy and procedure documentation lacks comprehensive coverage, regular updates, and enforcement mechanisms required for Level {target_level}."
        gap_initiative = "Develop comprehensive policy framework, establish regular review procedures, implement policy enforcement controls"
        common_gap = "Policy and Documentation Gap"
    elif any(keyword in evidence_lower for keyword in ['vendor', 'third', 'party', 'supplier']):
        gap_desc = f"Current vendor management lacks systematic due diligence, ongoing monitoring, and contract controls required for Level {target_level} compliance."
        gap_initiative = "Implement vendor risk assessment program, establish ongoing monitoring procedures, develop contract security requirements"
        common_gap = "Vendor Management Control Gap"
    elif any(keyword in evidence_lower for keyword in ['business', 'continuity', 'disaster']):
        gap_desc = f"Current business continuity planning lacks comprehensive testing, recovery procedures, and crisis management required for Level {target_level}."
        gap_initiative = "Develop business continuity plan, implement regular testing procedures, establish crisis management framework"
        common_gap = "Business Continuity Planning Gap"
    else:
        gap_desc = f"Current {domain.lower()} procedures lack documentation depth and control frameworks required for Level {target_level} {framework} compliance."
        gap_initiative = "Develop comprehensive procedures, implement process controls, establish compliance monitoring"
        common_gap = "Documentation and Process Control Gap"

    # Severity assessment based on maturity gap and domain
    if maturity_gap >= 3:
        severity = "HIGH - Critical compliance gap with immediate regulatory risks and potential business impact"
    elif maturity_gap >= 2:
        severity = "MEDIUM-HIGH - Significant compliance exposure requiring prioritized attention and resource allocation"
    elif maturity_gap >= 1:
        severity = "MEDIUM - Moderate gap requiring structured improvement and timeline commitment"
    else:
        severity = "LOW - Minor enhancement opportunity with minimal compliance risk"

    return {
        "gap_description": gap_desc,
        "gap_initiative": gap_initiative,
        "common_gap_description": common_gap,
        "severity_assessment": severity
    }

def parse_gap_analysis_result(result_text):
    """Parse the gap analysis result into components"""

    components = {
        "gap_description": "",
        "gap_initiative": "",
        "common_gap_description": "",
        "severity_assessment": ""
    }

    lines = result_text.split('\n')
    current_component = None

    for line in lines:
        line = line.strip()
        if line.startswith("1. Gap Description:") or line.startswith("Gap Description:"):
            current_component = "gap_description"
            components[current_component] = line.split(":", 1)[-1].strip()
        elif line.startswith("2. Gap Initiative:") or line.startswith("Gap Initiative:"):
            current_component = "gap_initiative"
            components[current_component] = line.split(":", 1)[-1].strip()
        elif line.startswith("3. Common Gap Description:") or line.startswith("Common Gap Description:"):
            current_component = "common_gap_description"
            components[current_component] = line.split(":", 1)[-1].strip()
        elif line.startswith("4. Severity Assessment:") or line.startswith("Severity Assessment:"):
            current_component = "severity_assessment"
            components[current_component] = line.split(":", 1)[-1].strip()
        elif current_component and line and not line.startswith(("1.", "2.", "3.", "4.")):
            components[current_component] += " " + line

    # If parsing fails, return structured fallback
    if not any(components.values()):
        return create_rule_based_gap_analysis(
            question="Parsing fallback",
            evidence=result_text,
            current_level=1,
            target_level=3,
            domain="General",
            framework="Standard"
        )

    return components

def create_maturity_visualization(current_level, target_level, confidence):
    """Create visualization for maturity levels"""

    # Gauge chart for current maturity
    gauge_fig = go.Figure(go.Indicator(
        mode = "gauge+number+delta",
        value = current_level,
        domain = {'x': [0, 1], 'y': [0, 1]},
        title = {'text': "Current Maturity Level"},
        delta = {'reference': target_level, 'position': "top"},
        gauge = {
            'axis': {'range': [None, 5]},
            'bar': {'color': "darkblue"},
            'steps': [
                {'range': [0, 1], 'color': "lightgray"},
                {'range': [1, 2], 'color': "gray"},
                {'range': [2, 3], 'color': "yellow"},
                {'range': [3, 4], 'color': "orange"},
                {'range': [4, 5], 'color': "green"}
            ],
            'threshold': {
                'line': {'color': "red", 'width': 4},
                'thickness': 0.75,
                'value': target_level
            }
        }
    ))

    gauge_fig.update_layout(height=300)

    return gauge_fig

def render_gap_analysis_form(gap_components):
    """Render editable gap analysis form"""

    st.markdown('<div class="gap-form">', unsafe_allow_html=True)
    st.subheader("🎯 Gap Analysis Results (Editable)")

    with st.form("gap_analysis_form"):
        # Gap Description
        st.markdown("**📋 Gap Description**")
        gap_description = st.text_area(
            "Gap Description:",
            value=gap_components.get("gap_description", ""),
            height=100,
            help="Describe what specific controls, processes, or documentation are missing",
            label_visibility="collapsed"
        )

        # Gap Initiative
        st.markdown("**🚀 Gap Initiative**")
        gap_initiative = st.text_area(
            "Gap Initiative:",
            value=gap_components.get("gap_initiative", ""),
            height=100,
            help="List the top 3 specific actions needed to close this gap",
            label_visibility="collapsed"
        )

        # Common Gap Description
        st.markdown("**🔍 Common Gap Category**")
        common_gap_description = st.text_input(
            "Common Gap Category:",
            value=gap_components.get("common_gap_description", ""),
            help="What is the general category of this compliance gap?",
            label_visibility="collapsed"
        )

        # Severity Assessment
        st.markdown("**⚠️ Severity Assessment**")
        severity_options = [
            "LOW - Minor enhancement opportunity with minimal compliance risk",
            "MEDIUM - Moderate gap requiring structured improvement and timeline commitment",
            "MEDIUM-HIGH - Significant compliance exposure requiring prioritized attention and resource allocation",
            "HIGH - Critical compliance gap with immediate regulatory risks and potential business impact"
        ]

        current_severity = gap_components.get("severity_assessment", "")
        default_severity_index = 0
        for i, option in enumerate(severity_options):
            if any(word in current_severity.upper() for word in option.split(" - ")[0].split("-")):
                default_severity_index = i
                break

        severity_assessment = st.selectbox(
            "Severity Level:",
            options=severity_options,
            index=default_severity_index,
            label_visibility="collapsed"
        )

        # Submit button
        submitted = st.form_submit_button("💾 Update Gap Analysis", type="primary")

        if submitted:
            updated_components = {
                "gap_description": gap_description,
                "gap_initiative": gap_initiative,
                "common_gap_description": common_gap_description,
                "severity_assessment": severity_assessment
            }
            st.session_state['updated_gap_components'] = updated_components
            st.success("✅ Gap analysis updated successfully!")
            st.rerun()

    st.markdown('</div>', unsafe_allow_html=True)

    # Return updated components if they exist in session state
    return st.session_state.get('updated_gap_components', gap_components)

def main():
    st.title("🎯 Compliance Maturity & Gap Analysis Platform")
    st.markdown("*Automated maturity assessment and gap analysis for regulatory compliance*")
    st.markdown("---")

    # Load models
    with st.spinner("🤖 Loading AI models from Hugging Face..."):
        maturity_model, maturity_tokenizer = load_maturity_model()
        gap_model, gap_tokenizer, gap_model_name = load_gap_analysis_model()
        evidence_model, evidence_tokenizer, evidence_types = load_evidence_recommendation_model()

    # Check model loading status
    models_loaded = True
    if maturity_model is None or maturity_tokenizer is None:
        st.error("❌ Failed to load maturity model.")
        models_loaded = False

    if gap_model is None or gap_tokenizer is None:
        st.error("❌ Failed to load gap analysis model.")
        models_loaded = False

    if evidence_model is None or evidence_tokenizer is None or evidence_types is None:
        st.warning("⚠️ Evidence recommendation model not loaded. Feature will be disabled.")
        evidence_model = None

    if not models_loaded:
        return

    if evidence_model is not None:
        st.success("✅ All AI models loaded successfully from Hugging Face!")
    else:
        st.info("ℹ️ Core models loaded. Evidence recommendation feature unavailable.")

    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")

        # Domain selection
        domain = st.selectbox(
            "Select Domain:",
            ["Banking", "Healthcare", "Technology", "Manufacturing", "General"],
            index=0
        )

        # Framework selection
        framework_options = {
            "Banking": ["RBI IT Framework", "Basel III", "PCI DSS"],
            "Healthcare": ["HIPAA", "HITECH", "FDA CFR"],
            "Technology": ["ISO 27001", "NIST", "SOC 2"],
            "Manufacturing": ["ISO 9001", "OSHA", "FDA QSR"],
            "General": ["ISO 27001", "NIST", "COBIT"]
        }

        framework = st.selectbox(
            "Select Framework:",
            framework_options.get(domain, ["ISO 27001"]),
            index=0
        )

        st.info("💡 Target maturity will be automatically set to Current + 1 level")

    # Main input section
    col1, col2 = st.columns([1, 1])

    with col1:
        st.subheader("📝 Question Input")
        
        # Create sub-columns for question and button
        question_col, button_col = st.columns([4, 1])
        
        with question_col:
            question = st.text_area(
                "Enter your compliance question:",
                placeholder="What documented procedures exist for user access provisioning?",
                height=120,
                help="Enter the specific compliance question you want to assess",
                key="main_question"
            )
        
        with button_col:
            st.markdown("<br>", unsafe_allow_html=True)  # Add some spacing
            
            # Evidence recommendation button - only enabled when question is entered and model is loaded
            recommend_button_disabled = not (question.strip() and evidence_model is not None)
            
            recommend_evidence = st.button(
                "💡 Get Evidence",
                disabled=recommend_button_disabled,
                help="Get AI recommendations for evidence types based on your question" if not recommend_button_disabled else "Enter a question first and ensure evidence model is loaded",
                type="secondary"
            )

    with col2:
        st.subheader("📄 Evidence Input")
        input_method = st.radio(
            "Choose input method:",
            ["Text Input", "File Upload"]
        )

    # Handle evidence recommendation
    if recommend_evidence and question.strip() and evidence_model is not None:
        with st.spinner("🔍 Analyzing question for evidence recommendations..."):
            recommendations = predict_evidence_types(
                evidence_model, 
                evidence_tokenizer, 
                evidence_types, 
                question, 
                domain, 
                framework,
                threshold=0.3,  # Lower threshold to show more options
                top_k=5
            )
            
            # Store recommendations in session state
            st.session_state['evidence_recommendations'] = recommendations
            st.session_state['evidence_question'] = question

    # Display evidence recommendations if they exist
    if 'evidence_recommendations' in st.session_state and 'evidence_question' in st.session_state:
        if st.session_state['evidence_question'] == question:  # Only show if question matches
            display_evidence_recommendations(st.session_state['evidence_recommendations'], question)
            
            # Add button to clear recommendations
            if st.button("❌ Clear Recommendations"):
                if 'evidence_recommendations' in st.session_state:
                    del st.session_state['evidence_recommendations']
                if 'evidence_question' in st.session_state:
                    del st.session_state['evidence_question']
                st.rerun()

    # Evidence input section
    evidence_text = ""

    if input_method == "Text Input":
        evidence_text = st.text_area(
            "Enter evidence text:",
            placeholder="Describe the current state of controls, documentation, or processes...",
            height=200,
            help="Provide detailed evidence about the current implementation"
        )
    else:
        uploaded_file = st.file_uploader(
            "Upload evidence file:",
            type=['txt', 'csv', 'xlsx', 'xls', 'pdf', 'docx'],
            help="Supported formats: TXT, CSV, XLSX, PDF, DOCX"
        )

        if uploaded_file is not None:
            with st.spinner("📄 Processing uploaded file..."):
                evidence_text = process_uploaded_file(uploaded_file)

            if evidence_text:
                file_size_kb = len(evidence_text.encode('utf-8')) / 1024
                st.success(f"✅ File uploaded successfully! ({len(evidence_text)} characters, {file_size_kb:.1f} KB)")

                # Enhanced preview with full document option
                col1_preview, col2_preview = st.columns([3, 1])

                with col1_preview:
                    preview_option = st.radio(
                        "Preview option:",
                        ["First 2000 characters", "Full document"],
                        horizontal=True
                    )

                with col2_preview:
                    if st.button("🔄 Refresh Preview"):
                        st.rerun()

                # Display preview based on selection
                with st.expander("📋 Document Preview", expanded=True):
                    if preview_option == "First 2000 characters":
                        preview_text = evidence_text[:2000]
                        if len(evidence_text) > 2000:
                            preview_text += "\n\n... (truncated, select 'Full document' to see complete content)"
                    else:
                        preview_text = evidence_text

                    # Use custom CSS class for better formatting
                    st.markdown(f'<div class="file-preview">{preview_text}</div>', unsafe_allow_html=True)

                    # Show document statistics
                    word_count = len(evidence_text.split())
                    line_count = len(evidence_text.split('\n'))

                    col1_stats, col2_stats, col3_stats = st.columns(3)
                    with col1_stats:
                        st.metric("📝 Characters", f"{len(evidence_text):,}")
                    with col2_stats:
                        st.metric("📄 Words", f"{word_count:,}")
                    with col3_stats:
                        st.metric("📋 Lines", f"{line_count:,}")

    # Analysis button
    st.markdown("---")
    analyze_button = st.button("🚀 Analyze Compliance Gap", type="primary", use_container_width=True)

    # Perform analysis
    if analyze_button:
        if not question.strip():
            st.error("❌ Please enter a question.")
        elif not evidence_text.strip():
            st.error("❌ Please provide evidence text.")
        else:
            with st.spinner("🔍 Analyzing compliance maturity and gaps..."):

                # Step 1: Predict current maturity
                maturity_input = prepare_maturity_input_text(question, evidence_text)
                current_maturity, confidence, all_probabilities = predict_maturity(
                    maturity_model, maturity_tokenizer, maturity_input
                )

                # Step 2: Set target maturity (current + 1, max 5)
                target_maturity = min(current_maturity + 1, 5)

                # Step 3: Generate gap analysis
                gap_prompt = create_gap_analysis_prompt(
                    question, evidence_text, current_maturity, target_maturity, domain, framework
                )

                gap_result = generate_gap_analysis(
                    gap_model, gap_tokenizer, gap_prompt, gap_model_name,
                    question, evidence_text, current_maturity, target_maturity, domain, framework
                )
                
                # Handle both string and dict returns from gap analysis
                if isinstance(gap_result, str):
                    gap_components = parse_gap_analysis_result(gap_result)
                else:
                    gap_components = gap_result

                # Calculate averages
                current_average = current_maturity
                target_average = target_maturity
                gap_size = target_maturity - current_maturity

            # Display results
            st.markdown("---")
            st.header("📊 Analysis Results")

            # Key metrics row
            col1, col2, col3, col4 = st.columns(4)

            with col1:
                st.metric(
                    label="Current Maturity",
                    value=f"Level {current_maturity}",
                    help=f"AI Confidence: {confidence:.2%}"
                )

            with col2:
                st.metric(
                    label="Target Maturity",
                    value=f"Level {target_maturity}",
                    delta=f"+{gap_size} levels" if gap_size > 0 else "At target"
                )

            with col3:
                st.metric(
                    label="Current Average",
                    value=f"{current_average:.1f}",
                    help="Current maturity level score"
                )

            with col4:
                st.metric(
                    label="Target Average",
                    value=f"{target_average:.1f}",
                    help="Target maturity level score"
                )

            # Visualization section
            st.subheader("📈 Maturity Level Visualization")

            col1, col2 = st.columns([1, 1])

            with col1:
                # Gauge chart
                gauge_fig = create_maturity_visualization(current_maturity, target_maturity, confidence)
                st.plotly_chart(gauge_fig, use_container_width=True)

            with col2:
                # Probability distribution
                prob_df = pd.DataFrame({
                    'Level': [f"Level {i}" for i in range(len(all_probabilities))],
                    'Probability': all_probabilities
                })

                bar_fig = px.bar(
                    prob_df,
                    x='Level',
                    y='Probability',
                    title="Maturity Level Confidence Distribution",
                    color='Probability',
                    color_continuous_scale='viridis'
                )
                st.plotly_chart(bar_fig, use_container_width=True)

            # Gap Analysis Form Section (Editable)
            final_gap_components = render_gap_analysis_form(gap_components)

            # Display severity with appropriate styling
            severity = final_gap_components.get("severity_assessment", "")
            if "HIGH" in severity.upper():
                st.error(f"🔴 {severity}")
            elif "MEDIUM" in severity.upper():
                st.warning(f"🟡 {severity}")
            else:
                st.info(f"🔵 {severity}")

            # Detailed recommendations
            st.subheader("💡 Detailed Recommendations")

            if gap_size > 0:
                recommendations = []

                if gap_size >= 3:
                    recommendations = [
                        "🎯 **Immediate Action Required** - This is a critical gap requiring executive attention",
                        "📋 Develop comprehensive transformation roadmap with clear milestones",
                        "👥 Assign dedicated project team with sufficient resources and budget",
                        "📊 Establish regular progress monitoring and reporting mechanisms",
                        "🔄 Implement phased approach with quick wins and long-term strategic initiatives"
                    ]
                elif gap_size >= 2:
                    recommendations = [
                        "⚡ **Prioritized Attention** - Significant improvement needed",
                        "📝 Create detailed implementation plan with defined timelines",
                        "🎓 Invest in training and capability development",
                        "🔧 Focus on process standardization and documentation",
                        "📈 Establish measurement and monitoring capabilities"
                    ]
                else:
                    recommendations = [
                        "🎯 **Focused Improvement** - Moderate gap requiring structured approach",
                        "📚 Enhance documentation and procedural clarity",
                        "🔍 Implement regular review and validation processes",
                        "📊 Establish basic measurement and reporting",
                        "🔄 Focus on continuous improvement practices"
                    ]

                for rec in recommendations:
                    st.markdown(f"• {rec}")
            else:
                st.success("🎉 **Target Achieved!** Current maturity meets or exceeds target level.")

            # Action Plan - Fixed array length issue
            st.subheader("📋 Suggested Action Plan")

            if gap_size > 0:
                # Create action plan with consistent array lengths
                num_phases = min(gap_size, 3)  # Limit to maximum 3 phases

                phases = [f'Phase {i+1}' for i in range(num_phases)]
                timelines = ['0-3 months', '3-6 months', '6-12 months'][:num_phases]
                focus_areas = [
                    'Foundation & Quick Wins',
                    'Process Implementation',
                    'Optimization & Maturity'
                ][:num_phases]
                key_activities = [
                    'Documentation, Training, Basic Controls',
                    'Process Automation, Monitoring, Compliance',
                    'Continuous Improvement, Advanced Analytics'
                ][:num_phases]

                action_plan = pd.DataFrame({
                    'Phase': phases,
                    'Timeline': timelines,
                    'Focus Area': focus_areas,
                    'Key Activities': key_activities
                })

                st.dataframe(action_plan, use_container_width=True)
            else:
                st.info("🎯 No action plan needed - current maturity level meets target.")

            # Export functionality
            st.subheader("💾 Export Results")

            # Create summary report
            summary_report = f"""
# Compliance Gap Analysis Report

**Assessment Date:** {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}
**Domain:** {domain}
**Framework:** {framework}

## Question Analyzed
{question}

## Evidence Provided
{evidence_text[:20000]}{'...' if len(evidence_text) > 20000 else ''}

## Maturity Assessment
- **Current Maturity Level:** {current_maturity} ({MATURITY_LEVELS[current_maturity]})
- **Target Maturity Level:** {target_maturity} ({MATURITY_LEVELS[target_maturity]})
- **AI Confidence:** {confidence:.2%}
- **Gap Size:** {gap_size} level(s)

## Gap Analysis
**Gap Description:** {final_gap_components.get("gap_description", "N/A")}

**Gap Initiative:** {final_gap_components.get("gap_initiative", "N/A")}

**Common Gap Category:** {final_gap_components.get("common_gap_description", "N/A")}

**Severity Assessment:** {final_gap_components.get("severity_assessment", "N/A")}

## Recommendations
{chr(10).join(f"• {rec.replace('*', '').replace('🎯', '').replace('⚡', '').replace('🎯', '').strip()}" for rec in recommendations) if 'recommendations' in locals() else "No specific recommendations generated."}

## Action Plan
{action_plan.to_string(index=False) if 'action_plan' in locals() and not action_plan.empty else "No action plan required."}
            """

            col1, col2 = st.columns([1, 1])

            with col1:
                st.download_button(
                    label="📄 Download Summary Report",
                    data=summary_report,
                    file_name=f"gap_analysis_report_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.md",
                    mime="text/markdown"
                )

            with col2:
                # Export as CSV
                results_df = pd.DataFrame({
                    'Metric': ['Current Maturity', 'Target Maturity', 'Gap Size', 'Confidence', 'Domain', 'Framework'],
                    'Value': [current_maturity, target_maturity, gap_size, f"{confidence:.2%}", domain, framework]
                })

                csv_data = results_df.to_csv(index=False)
                st.download_button(
                    label="📊 Download Data (CSV)",
                    data=csv_data,
                    file_name=f"gap_analysis_data_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv"
                )

            # Store results in session state for persistence
            st.session_state['analysis_results'] = {
                'current_maturity': current_maturity,
                'target_maturity': target_maturity,
                'confidence': confidence,
                'gap_components': final_gap_components,
                'question': question,
                'evidence': evidence_text[:20000],
                'domain': domain,
                'framework': framework
            }

    # Add section to display previous results if they exist (only when no analysis is running)
    if 'analysis_results' in st.session_state and not analyze_button and question.strip() == "" and evidence_text.strip() == "":
        st.markdown("---")
        st.subheader("📋 Previous Analysis Results")

        prev_results = st.session_state['analysis_results']

        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Previous Current Maturity", f"Level {prev_results['current_maturity']}")
        with col2:
            st.metric("Previous Target Maturity", f"Level {prev_results['target_maturity']}")
        with col3:
            st.metric("Previous Confidence", f"{prev_results['confidence']:.2%}")

        with st.expander("View Previous Gap Analysis"):
            st.write("**Gap Description:**", prev_results['gap_components'].get('gap_description', 'N/A'))
            st.write("**Gap Initiative:**", prev_results['gap_components'].get('gap_initiative', 'N/A'))
            st.write("**Common Gap Category:**", prev_results['gap_components'].get('common_gap_description', 'N/A'))
            st.write("**Severity:**", prev_results['gap_components'].get('severity_assessment', 'N/A'))

        # Add button to clear previous results
        if st.button("🗑️ Clear Previous Results"):
            if 'analysis_results' in st.session_state:
                del st.session_state['analysis_results']
            if 'updated_gap_components' in st.session_state:
                del st.session_state['updated_gap_components']
            if 'evidence_recommendations' in st.session_state:
                del st.session_state['evidence_recommendations']
            if 'evidence_question' in st.session_state:
                del st.session_state['evidence_question']
            st.rerun()

if __name__ == "__main__":
    main()
